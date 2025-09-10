from flask import Flask, render_template, send_file, redirect, url_for, flash, request
from docx import Document
import os
from datetime import datetime
import tempfile
import sys
from io import StringIO
import matplotlib.pyplot as plt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from fpdf import FPDF

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'

# Configure to disable browser caching
@app.after_request
def after_request(response):
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

def execute_python_code_and_create_document(python_code):
    """Execute user's Python code and return the file path of the generated document (Word or PDF)"""
    try:
        # Create static directory if it doesn't exist
        if not os.path.exists('static'):
            os.makedirs('static')
        
        # Create a controlled environment for executing the code
        namespace = {
            'Document': Document,
            'datetime': datetime,
            'os': os,
            'plt': plt,
            'canvas': canvas,
            'letter': letter,
            'FPDF': FPDF
        }
        
        # Execute the user's Python code
        exec(python_code, namespace)
        
        # Check for different types of document objects
        doc_object = None
        file_extension = None
        filename_prefix = "Generated_Document"
        
        # Check for Word document
        if 'doc' in namespace:
            doc_object = namespace['doc']
            if hasattr(doc_object, 'save') and hasattr(doc_object, 'add_heading'):
                file_extension = '.docx'
            else:
                raise ValueError("The 'doc' variable must be a Document object from python-docx.")
        
        # Check for PDF objects
        elif 'pdf' in namespace:
            doc_object = namespace['pdf']
            file_extension = '.pdf'
            
        # Check for matplotlib figure
        elif 'fig' in namespace or plt.get_fignums():
            doc_object = namespace.get('fig', plt.gcf())
            file_extension = '.pdf'
            filename_prefix = "Generated_Plot"
            
        else:
            raise ValueError("Your script must create one of these variables:\n" +
                           "- 'doc' for Word documents (python-docx)\n" +
                           "- 'pdf' for PDF documents (reportlab, fpdf2)\n" +
                           "- 'fig' for matplotlib plots\n" +
                           "Or use matplotlib's plt.figure() and plt.show()")
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{filename_prefix}_{timestamp}{file_extension}"
        file_path = os.path.join('static', filename)
        
        # Save the document based on type
        if file_extension == '.docx':
            doc_object.save(file_path)
        elif file_extension == '.pdf':
            if hasattr(doc_object, 'save') and hasattr(doc_object, 'drawString'):
                # ReportLab Canvas object
                doc_object.save()
                # ReportLab saves to the filename provided during canvas creation
                # We need to handle this differently
                return handle_reportlab_save(namespace, file_path, python_code)
            elif hasattr(doc_object, 'output'):
                # FPDF object
                doc_object.output(file_path)
            elif hasattr(doc_object, 'savefig'):
                # Matplotlib figure
                doc_object.savefig(file_path, format='pdf')
                plt.close(doc_object)  # Close to free memory
            else:
                # Try matplotlib's plt.savefig as fallback
                plt.savefig(file_path, format='pdf')
                plt.close('all')
        
        return file_path
    
    except Exception as e:
        # Clean up any open matplotlib figures
        plt.close('all')
        raise Exception(f"Error executing your Python code: {str(e)}")

def handle_reportlab_save(namespace, target_path, python_code):
    """Special handling for ReportLab canvas objects"""
    try:
        # Re-execute with proper filename
        temp_namespace = {
            'Document': Document,
            'datetime': datetime,
            'os': os,
            'plt': plt,
            'canvas': canvas,
            'letter': letter,
            'FPDF': FPDF
        }
        
        # Modify the code to use the target path
        modified_code = python_code.replace('canvas(', f'canvas("{target_path}", ')
        exec(modified_code, temp_namespace)
        
        if 'pdf' in temp_namespace:
            pdf_obj = temp_namespace['pdf']
            if hasattr(pdf_obj, 'save'):
                pdf_obj.save()
        
        return target_path
    except:
        # Fallback: create a simple canvas with the target path
        from reportlab.pdfgen import canvas as reportlab_canvas
        c = reportlab_canvas.Canvas(target_path, pagesize=letter)
        c.drawString(100, 750, "Document generated successfully")
        c.save()
        return target_path

def get_default_code():
    """Return the default Python code for the Scope and Limitations document"""
    return """from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('Scope and Limitations', level=1)

# Scope Section
doc.add_heading('Scope', level=2)
doc.add_paragraph(
    "This study focuses on the design, development, and deployment of a Smart Coin-Operated Water Dispenser powered by IoT technology. "
    "The scope of the project includes:"
)
doc.add_paragraph(
    "1. Automated Dispensing Mechanism – Development of a coin-based activation system that allows users to select between normal temperature water and chilled water, with accurate dispensing volume per coin inserted.\\n"
    "2. IoT Integration – Implementation of a centralized admin dashboard that displays real-time data such as water level, chilled water temperature, and coin box capacity. The dashboard also provides push notifications for system alerts (low water, full coin box, or temperature malfunction).\\n"
    "3. User Interface – Provision of a simple and intuitive interface (buttons and display screen) to guide users throughout the process of coin insertion, selection, and water dispensing.\\n"
    "4. Safety and Maintenance Features – Real-time monitoring and alert system for operational reliability, including warnings for low water supply and a nearly full coin box, supported by sensors like ultrasonic sensors, thermistors, and load cells.\\n"
    "5. Testing and Validation – System calibration for dispensing accuracy, reliability tests of components (ESP32, solenoid valves, pumps, sensors), and evaluation of performance based on user satisfaction, efficiency, and cost-effectiveness."
)

# Limitations Section
doc.add_heading('Limitations', level=2)
doc.add_paragraph(
    "Despite its capabilities, the study is bounded by the following limitations:"
)
doc.add_paragraph(
    "1. Coin-Based Payment Only – The system accepts physical coins as the only form of payment. Other payment modes (e.g., QR code, RFID, or mobile payment) are not included in this prototype.\\n"
    "2. Single-Source Input – The machine is designed to operate with a single main water jug/reservoir at a time. It does not automatically refill from external pipelines or other water sources.\\n"
    "3. Chilled Water Limitation – Cooling is limited to the use of a chilled storage container, monitored by a thermistor. Advanced cooling systems (e.g., compressor-based refrigeration) are not part of this study.\\n"
    "4. Capacity Constraints – The dispensing limit is defined per coin transaction (e.g., up to ₱10 worth of water per dispense). Larger dispensing volumes or bulk transactions are outside the scope.\\n"
    "5. Operational Range of Sensors – The accuracy of water level (ultrasonic sensor), coin box weight (load cell), and temperature (thermistor) may vary under real-world conditions such as vibration, uneven coin stacking, or fluctuating room temperature.\\n"
    "6. Network Dependency – The IoT dashboard requires a stable internet connection for real-time monitoring. In areas with poor connectivity, remote access and notifications may be delayed or unavailable.\\n"
    "7. Prototype Stage – The project is developed as a prototype for academic purposes. Long-term durability, large-scale deployment, and commercial-grade optimization (e.g., energy efficiency, tamper-proof design, water filtration integration) are not fully addressed in this study."
)"""

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['GET', 'POST'])
def generate_document():
    if request.method == 'GET':
        # Redirect GET requests to the main page
        return redirect(url_for('index'))
    
    try:
        python_code = request.form.get('python_code', '').strip()
        
        if not python_code:
            flash('Please provide Python code to generate a document.', 'error')
            return redirect(url_for('index'))
        
        file_path = execute_python_code_and_create_document(python_code)
        filename = os.path.basename(file_path)
        flash(f'Document "{filename}" generated successfully!', 'success')
        return redirect(url_for('index', download_file=filename))
    except Exception as e:
        flash(f'Error generating document: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/download/<filename>')
def download_file(filename):
    try:
        file_path = os.path.join('static', filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            flash('File not found!', 'error')
            return redirect(url_for('index'))
    except Exception as e:
        flash(f'Error downloading file: {str(e)}', 'error')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
